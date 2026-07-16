"""
Motor AutoOferta (versión de despliegue, todo en un módulo).

De un PDF de pliego a: criterios extraídos + borrador de memoria técnica (docx) + checklist.
Necesita ANTHROPIC_API_KEY para los pasos de IA. El modelo se ajusta con ANTHROPIC_MODEL.
"""
import os
import json
from pathlib import Path

import pdfplumber
import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

DEFAULT_MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-5")
_client = None


def _get_client():
    global _client
    if _client is None:
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            raise RuntimeError("Falta ANTHROPIC_API_KEY.")
        _client = anthropic.Anthropic(api_key=api_key)
    return _client


def _complete_json(system, user, max_tokens=4000, model=None):
    client = _get_client()
    msg = client.messages.create(
        model=model or DEFAULT_MODEL, max_tokens=max_tokens,
        system=system, messages=[{"role": "user", "content": user}],
    )
    texto = "".join(b.text for b in msg.content if b.type == "text").strip()
    if texto.startswith("```"):
        texto = texto.split("```", 2)[1]
        if texto.startswith("json"):
            texto = texto[4:]
        texto = texto.strip().rstrip("`").strip()
    return json.loads(texto)


# ---------------- Paso 0: PDF -> texto ----------------
def extract_text_from_pdf(pdf_path):
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"No existe el pliego: {pdf_path}")
    partes = []
    with pdfplumber.open(pdf_path) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            t = page.extract_text() or ""
            if t.strip():
                partes.append(f"\n\n----- PÁGINA {i} -----\n{t}")
    texto = "\n".join(partes).strip()
    if not texto:
        raise ValueError("No se pudo extraer texto (¿PDF escaneado? haría falta OCR).")
    return texto


def clean_text(texto, max_chars=120_000):
    limpio = "\n".join(l.rstrip() for l in texto.splitlines() if l.strip())
    if len(limpio) > max_chars:
        limpio = limpio[:max_chars] + "\n\n[... texto recortado por longitud ...]"
    return limpio


# ---------------- Prompts ----------------
ANALYZE_SYSTEM = """Eres un experto en contratación pública española con 15 años preparando
ofertas ganadoras para pymes. Lee un pliego (PCAP y/o PPT) y extrae con precisión la información
que una empresa necesita para decidir si se presenta y preparar su oferta.
Reglas: no inventes nada (dato ausente -> null y a "faltantes"); cita cifras y puntuaciones tal
cual; los criterios de adjudicación con su puntuación máxima y si son por fórmula o juicio de valor
son lo más importante. Devuelve EXCLUSIVAMENTE un objeto JSON válido, sin texto alrededor."""

ANALYZE_USER = """Analiza este pliego y devuelve un JSON con esta forma exacta:
{{
  "objeto": "...", "organo_contratacion": "... o null", "presupuesto_base": "... o null",
  "valor_estimado": "... o null", "duracion": "... o null", "plazo_presentacion": "... o null",
  "criterios_adjudicacion": [{{"nombre":"...","puntos_max":60,"tipo":"formula|juicio_valor","detalle":"..."}}],
  "solvencia_economica": "... o null", "solvencia_tecnica": "... o null",
  "documentacion": [{{"documento":"...","sobre":"1|2|3","obligatorio":true}}],
  "exclusiones_criticas": ["..."], "faltantes": ["..."]
}}

TEXTO DEL PLIEGO:
\"\"\"
{pliego}
\"\"\""""

DRAFT_SYSTEM = """Eres un consultor senior que redacta memorias técnicas GANADORAS para
licitaciones públicas en España, para pymes reales.
Principios: estructura la memoria según los criterios de juicio de valor, dando más extensión a
los que más puntúan; usa SOLO datos reales del perfil (dato ausente -> marcador entre corchetes
como [INDICAR ...], nunca te lo inventes); tono profesional y concreto; es un BORRADOR de
asistencia que la empresa revisa y firma."""

DRAFT_USER = """Redacta un borrador de MEMORIA TÉCNICA para esta licitación, estructurado según sus
criterios de juicio de valor y optimizado para puntuar. Devuelve un JSON con esta forma:
{{
  "titulo":"Memoria Técnica — <objeto>",
  "resumen_ejecutivo":"...",
  "secciones":[{{"titulo":"...","puntos_asociados":20,"contenido":"..."}}],
  "avisos_revision":["..."]
}}

ANÁLISIS DEL PLIEGO:
{analisis}

PERFIL DE LA EMPRESA:
{perfil}"""


# ---------------- Paso 1: analizar ----------------
def analyze_pliego(texto_pliego):
    return _complete_json(ANALYZE_SYSTEM, ANALYZE_USER.format(pliego=texto_pliego), 4000)


# ---------------- Paso 2: redactar ----------------
def draft_memoria(analisis, perfil):
    user = DRAFT_USER.format(
        analisis=json.dumps(analisis, ensure_ascii=False, indent=2),
        perfil=json.dumps(perfil, ensure_ascii=False, indent=2),
    )
    return _complete_json(DRAFT_SYSTEM, user, 6000)


# ---------------- Paso 3: checklist ----------------
def build_checklist(analisis):
    docs = analisis.get("documentacion") or []
    por_sobre = {}
    for d in docs:
        sobre = str(d.get("sobre") or "Sin especificar")
        por_sobre.setdefault(sobre, []).append({
            "documento": d.get("documento", ""),
            "obligatorio": bool(d.get("obligatorio", True)), "hecho": False,
        })
    avisos = []
    if analisis.get("plazo_presentacion"):
        avisos.append(f"Plazo de presentación: {analisis['plazo_presentacion']}")
    for exc in (analisis.get("exclusiones_criticas") or []):
        avisos.append(f"⚠ Exclusión directa: {exc}")
    for falta in (analisis.get("faltantes") or []):
        avisos.append(f"Revisar en el pliego original: {falta}")
    return {"checklist_por_sobre": por_sobre,
            "solvencia": {"economica": analisis.get("solvencia_economica"),
                          "tecnica": analisis.get("solvencia_tecnica")},
            "avisos": avisos}


# ---------------- Render docx ----------------
def render_memoria_docx(memoria, analisis, salida):
    salida = Path(salida)
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    t = doc.add_heading(memoria.get("titulo", "Memoria Técnica"), level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if analisis.get("organo_contratacion"):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Órgano de contratación: {analisis['organo_contratacion']}"); r.italic = True
    a = doc.add_paragraph(); a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = a.add_run("BORRADOR generado por AutoOferta · Revíselo, complételo y fírmelo antes de presentar.")
    r.font.color.rgb = RGBColor(0xB0, 0x8A, 0x00); r.bold = True
    doc.add_paragraph()
    if memoria.get("resumen_ejecutivo"):
        doc.add_heading("Resumen ejecutivo", level=1)
        for par in memoria["resumen_ejecutivo"].split("\n"):
            if par.strip():
                doc.add_paragraph(par.strip())
    for sec in memoria.get("secciones", []):
        pts = sec.get("puntos_asociados")
        enc = sec.get("titulo", "Sección") + (f"  ·  {pts} puntos" if pts else "")
        doc.add_heading(enc, level=1)
        for par in str(sec.get("contenido", "")).split("\n"):
            if par.strip():
                doc.add_paragraph(par.strip())
    avisos = memoria.get("avisos_revision") or []
    if avisos:
        doc.add_page_break(); doc.add_heading("Antes de presentar: revise estos puntos", level=1)
        for x in avisos:
            doc.add_paragraph(x, style="List Bullet")
    doc.save(salida)
    return salida


# ---------------- Orquestador ----------------
def procesar_licitacion(pliego_pdf, perfil, carpeta_salida):
    carpeta = Path(carpeta_salida); carpeta.mkdir(parents=True, exist_ok=True)
    texto = clean_text(extract_text_from_pdf(pliego_pdf))
    analisis = analyze_pliego(texto)
    (carpeta / "analisis.json").write_text(json.dumps(analisis, ensure_ascii=False, indent=2), encoding="utf-8")
    memoria = draft_memoria(analisis, perfil)
    docx_path = render_memoria_docx(memoria, analisis, carpeta / "memoria_tecnica.docx")
    checklist = build_checklist(analisis)
    (carpeta / "checklist.json").write_text(json.dumps(checklist, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"analisis": analisis, "memoria": memoria, "checklist": checklist,
            "memoria_docx": str(docx_path)}
